# coding:utf-8
# import docx
# # from win32com import client as wc
# import win32com
# # 首先将doc转换成docx
# word = wc.Dispatch("Word.Application")
#
# doc = word.Documents.Open(r"D:\\demo.doc")
# #使用参数16表示将doc转换成docx
# doc.SaveAs(r"D:\\most.docx",16)
# doc.Close()
#
# word.Quit()
#
# #读取word内容
# doc = docx.Document("D:\most.docx")
# data = doc.paragraphs[0].text
# print(data)

# 转换成html
# from docx2html import convert
# import HTMLParser
#
# html_parser = HTMLParser.HTMLParser()
# #使用docx2html模块将docx文件转成html串，随后你想干嘛都行
# html = convert("G:\\t.docx")
#
# # 这句非常关键，docx2html模块将中文进行了转义，所以要将生成的字符串重新转义
# print html_parser.enescape(html)

import mammoth
import mammoth.transforms
import os
from docx import Document
from bson import json_util
import zipfile
import json
import unidecode
import requests

guidUrl = "https://my.phrplus.com/REST/guid"
inputPath = '/Users/admin/cwell/parser/docfiles/'
imgPath = "/Users/admin/cwell/parser/imgs/"
outputFile = '/Users/admin/cwell/parser/output/output.json'

styleMap = """
p[style-name='Title'] => h1.hide
p[style-name='Subhead 1'] => h3
p[style-name='List Bullet'] => ul.first > li:fresh
p[style-name='List Bullet 2'] => ul.second > li:fresh
p[style-name='Hyperlink']=>a.link
"""


def convert_image(image):
    return {
        "src": ""
    }


def parseFile():
    document = Document('tt.docx')
    article = {"Title": document.core_properties.title, "Content": ""}
    with open('tt.docx', "rb") as docFile:
        html = mammoth.convert_to_html(docFile, style_map=styleMap)
        decoded = unidecode.unidecode(html.value)

    if not article["Title"]:
        for para in document.paragraphs:
            if para.style.name == 'Title':
                if para.text:
                    article["Title"] = para.text

    article["Content"] = decoded
    return article

with open('ttt.html', 'wb') as f:
    f.write(parseFile()["Content"])